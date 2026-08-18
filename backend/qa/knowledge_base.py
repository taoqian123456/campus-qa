"""
知识库文档处理：文本提取、分块。
"""
import re
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from config import CHUNK_SIZE, CHUNK_OVERLAP, OCR_ENABLED


def extract_text(file_path: str | Path) -> str:
    """提取文档全文。支持 PDF（pypdf，扫描版自动走 OCR）、TXT、DOCX（python-docx，含表格），其他格式抛异常。"""
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        text = _extract_pdf_text(file_path)
        # 扫描版 PDF 没有文本层：OCR 兜底（每个文件只初始化一次模型）
        if not text.strip() and OCR_ENABLED:
            text = _extract_pdf_ocr(file_path)
        return text

    if suffix == ".txt":
        return file_path.read_text(encoding="utf-8")

    if suffix == ".docx":
        return _extract_docx(file_path)

    raise ValueError(f"暂不支持的文档格式: {suffix}")


def _extract_pdf_text(file_path: Path) -> str:
    """pypdf 提取 PDF 文本层；文件损坏/加密时抛 ValueError。"""
    try:
        reader = PdfReader(str(file_path))
    except Exception as e:
        raise ValueError(f"PDF 文件无法打开（可能已损坏或已加密）：{file_path.name}（{e}）")
    return "\n".join(page.extract_text() or "" for page in reader.pages)


# 全局单例，避免 rebuild 时每个扫描版 PDF 都重新加载一遍模型
_ocr = None


def _extract_pdf_ocr(file_path: Path) -> str:
    """OCR 兜底：把 PDF 每页渲染成图片，用 RapidOCR（PaddleOCR 中文模型）识别。

    后处理策略：
    - 页眉页脚残片（纯数字行、短日期行）识别前过滤，避免生成无信息量的垃圾块；
    - 页内 OCR 行用换行拼接（OCR 常把句子拦腰拆行，换行不会被分块器当成段落边界）；
    - 页与页之间也只用换行连接，跨页内容自然衔接，不再产生「【翻页】」独立段。
    """
    global _ocr
    if _ocr is None:
        # 依赖较重，延迟导入：不用 OCR 的部署不装这两个包也能跑
        from rapidocr_onnxruntime import RapidOCR
        _ocr = RapidOCR()

    import pypdfium2  # 轻量渲染库，同样延迟导入
    try:
        pdf = pypdfium2.PdfDocument(str(file_path))
    except Exception as e:
        raise ValueError(f"PDF 无法渲染（可能已损坏）：{file_path.name}（{e}）")

    pages: list[str] = []
    try:
        import numpy as np  # PIL Image 转 numpy 数组（rapidocr 1.x 不直接接受 PIL Image）
        for i in range(len(pdf)):
            bitmap = pdf[i].render(scale=2.0)  # 2 倍放大提高小字识别率
            pil_image = bitmap.to_pil()
            try:
                result = _ocr(np.array(pil_image))
            finally:
                pil_image.close()
                bitmap.close()
            # rapidocr 1.x 返回 (结果, 耗时)；每条结果 = [box, text, score]
            items = result[0] if isinstance(result, tuple) else result
            lines = [str(item[1]).strip() for item in (items or []) if len(item) > 1 and str(item[1]).strip()]
            pages.append("\n".join(l for l in lines if _is_content_line(l)))
    finally:
        pdf.close()
    # 页内 OCR 行用换行拼接（修复被拆行的句子）；页间留一个空行：
    # chunk_text 按空行切硬段，硬段内句子不跨页 => 每页（或每页一段）成为一个聚焦块，检索分数更高
    return "\n\n".join(p for p in pages if p.strip())


def _is_content_line(line: str) -> bool:
    """过滤 OCR 噪声行：页码（- 2 -、2 / 293）、页脚日期残片（2025年7月,日）等无信息量行。

    只丢「不含汉字」或「短日期残片」的行，政策正文（含数字条款如「1～3 学期」）不会误伤。
    """
    if not line:
        return False
    if not re.search(r"[一-鿿]", line):
        # 不含汉字：几乎不可能是政策正文，纯数字/标点行视为页眉页脚残片
        return len(line) >= 20
    # 短日期残片：只含数字、日期字（年/月/日）、标点（含英文句点），几乎必是落款/印发日期被 OCR 拆出来的碎片
    if len(line) < 12 and re.fullmatch(r"[\d年月日。，,;；:.\s·\-–—/]*", line):
        return False
    return True


def _extract_docx(file_path: Path) -> str:
    """提取 DOCX 全文：正文段落 + 表格（政策文件里表格常承载关键规则，一并拼入）。"""
    try:
        doc = DocxDocument(str(file_path))
    except Exception as e:
        raise ValueError(f"DOCX 文件无法打开（可能已损坏）：{file_path.name}（{e}）")

    parts: list[str] = []

    # 正文段落：跳过空段落（保留顺序）
    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    parts.extend(paras)

    # 表格：逐单元格提取，跳过空单元格；表格前插空行使其独立成段
    # （chunk_text 按空行切硬段，表格独立成段避免与正文粘连，检索时更容易整表命中）
    for ti, table in enumerate(doc.tables, 1):
        cells = []
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in cells:  # 合并单元格会重复返回，去重
                    cells.append(text)
        if cells:
            parts.append("")  # 与上文之间留空行
            parts.append(f"【表格{ti}】" + "；".join(cells))

    # 保留空字符串作为空行分隔符（chunk_text 按空行切硬段，表格独立成段）
    return "\n".join(parts)


def chunk_text(text: str, chunk_size: int = None, overlap: int = None) -> list[str]:
    """按段落把全文切成块，每块最多 chunk_size 字符，相邻块重叠约 overlap 字符。

    规则：
    - 先按空行切成「硬段」，硬段不会被打散；
    - 硬段内部按句子边界滚窗：窗口 = 上一块的 overlap 尾巴（取整句）+ 若干整句，
      保证句子永远不会被块边界拦腰截断；
    - 单句超过 chunk_size 时按字符截断滚动，相邻块共享 overlap 字符。
    """
    chunk_size = chunk_size or CHUNK_SIZE
    overlap = overlap if overlap is not None else CHUNK_OVERLAP

    if chunk_size <= overlap:
        raise ValueError(f"chunk_size({chunk_size}) 必须大于 overlap({overlap})")
    if chunk_size <= 0:
        raise ValueError(f"chunk_size 必须为正数: {chunk_size}")

    hard_paras = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: list[str] = []
    carry = ""  # 上一块留给下一块的 overlap 尾巴

    def flush(window: list[str]) -> str:
        """落块；返回下一块的 overlap 尾巴（取窗口尾部整句，总长不超过 overlap）。"""
        chunks.append("".join(window))
        tail = ""
        for s in reversed(window):
            if len(tail) + len(s) <= overlap:
                tail = s + tail
            else:
                break
        return tail

    for para in hard_paras:
        sentences = re.split(r"(?<=[。！？!?；;])\s*", para)
        sentences = [s.strip() for s in sentences if s.strip()]
        if not sentences:
            continue

        window: list[str] = [carry] if carry else []
        window_chars = len(carry)

        for s in sentences:
            if window_chars + len(s) <= chunk_size:
                window.append(s)
                window_chars += len(s)
                continue

            # 放不下：先落块，取整句尾巴作为下一块的 overlap
            if window:
                carry = flush(window)
                window = [carry] if carry else []
                window_chars = len(carry)

            # 单句超长：按字符截断滚动，每块吃掉 fit 个新字符，其中 overlap 个在下一块重复
            while window_chars + len(s) > chunk_size:
                fit = chunk_size - window_chars
                if fit <= 0:  # 防御：尾巴占满整块（正常不会发生）
                    chunks.append("".join(window))
                    window = []
                    window_chars = 0
                    continue
                window.append(s[:fit])
                chunks.append("".join(window))
                skip = fit - overlap
                s = s[fit:] if skip <= 0 else s[skip:]  # skip<=0 时丢弃重叠，保证有进展
                window = []
                window_chars = 0
            if s:
                window.append(s)
                window_chars += len(s)

        if window_chars > 0:
            carry = flush(window)
        else:
            carry = ""

    return chunks
